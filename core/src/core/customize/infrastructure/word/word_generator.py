"""
Report Builder - Infrastructure Layer - Word

Word文書生成の実装（python-docxを利用）
"""

from __future__ import annotations

import tempfile
import uuid
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path

from core.logging_helper import get_logger
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

logger = get_logger("WordGenerator")


@dataclass
class ReportSectionContent:
    """レポートセクションのコンテンツ"""

    heading: str
    question: str
    content: str
    chart_paths: list[str] = field(default_factory=list)


class WordGenerator:
    """Word文書生成を担当するインフラストラクチャ

    python-docxを使用してWord文書を生成する。
    """

    def __init__(self) -> None:
        self._doc: Document | None = None

    def generate(
        self,
        title: str,
        summary: str,
        sections: list[ReportSectionContent],
        conclusion: str,
        author: str | None = None,
        created_at: datetime | None = None,
    ) -> str:
        """Word文書を生成し、ローカルパスを返す

        Args:
            title: レポートタイトル
            summary: エグゼクティブサマリー
            sections: 各セクションのコンテンツ
            conclusion: 結論
            author: 作成者
            created_at: 作成日時

        Returns:
            生成されたWordファイルのローカルパス
        """
        logger.info(f"Generating Word document: {title}")

        self._doc = Document()
        self._setup_styles()

        # タイトルページ
        self._add_title_page(title, author, created_at)

        # 目次（簡易版）
        self._add_table_of_contents(sections)

        # エグゼクティブサマリー
        self._add_summary_section(summary)

        # 各セクション
        for i, section in enumerate(sections, 1):
            self._add_analysis_section(i, section)

        # 結論
        self._add_conclusion_section(conclusion)

        # 保存
        local_path = self._save()
        logger.info(f"Word document generated: {local_path}")

        return local_path

    def _setup_styles(self) -> None:
        """スタイルを設定"""
        if self._doc is None:
            return

        styles = self._doc.styles

        # 見出し1のスタイル調整
        if "Heading 1" in [s.name for s in styles]:
            heading1 = styles["Heading 1"]
            heading1.font.size = Pt(18)
            heading1.font.bold = True
            heading1.font.color.rgb = RGBColor(0, 51, 102)

        # 見出し2のスタイル調整
        if "Heading 2" in [s.name for s in styles]:
            heading2 = styles["Heading 2"]
            heading2.font.size = Pt(14)
            heading2.font.bold = True
            heading2.font.color.rgb = RGBColor(0, 76, 153)

    def _add_title_page(
        self,
        title: str,
        author: str | None,
        created_at: datetime | None,
    ) -> None:
        """タイトルページを追加"""
        if self._doc is None:
            return

        # タイトル
        title_para = self._doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 51, 102)

        # スペース
        self._doc.add_paragraph()
        self._doc.add_paragraph()

        # サブタイトル
        subtitle_para = self._doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_para.add_run("データ分析レポート")
        subtitle_run.font.size = Pt(16)
        subtitle_run.font.color.rgb = RGBColor(102, 102, 102)

        # スペース
        for _ in range(5):
            self._doc.add_paragraph()

        # 作成者
        if author:
            author_para = self._doc.add_paragraph()
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_run = author_para.add_run(f"作成者: {author}")
            author_run.font.size = Pt(12)

        # 作成日
        date_str = (created_at or datetime.now()).strftime("%Y年%m月%d日")
        date_para = self._doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(f"作成日: {date_str}")
        date_run.font.size = Pt(12)

        # 改ページ
        self._doc.add_page_break()

    def _add_table_of_contents(self, sections: list[ReportSectionContent]) -> None:
        """目次を追加（簡易版）"""
        if self._doc is None:
            return

        self._doc.add_heading("目次", level=1)

        # 簡易目次
        toc_items = [
            "1. エグゼクティブサマリー",
        ]
        for i, section in enumerate(sections, 2):
            toc_items.append(f"{i}. {section.heading}")
        toc_items.append(f"{len(sections) + 2}. 結論")

        for item in toc_items:
            para = self._doc.add_paragraph(item)
            para.paragraph_format.left_indent = Inches(0.5)

        self._doc.add_page_break()

    def _add_summary_section(self, summary: str) -> None:
        """エグゼクティブサマリーを追加"""
        if self._doc is None:
            return

        self._doc.add_heading("エグゼクティブサマリー", level=1)
        self._doc.add_paragraph(summary)
        self._doc.add_paragraph()

    def _add_analysis_section(self, index: int, section: ReportSectionContent) -> None:
        """分析セクションを追加"""
        if self._doc is None:
            return

        # セクション見出し
        self._doc.add_heading(section.heading, level=1)

        # 質問
        question_para = self._doc.add_paragraph()
        question_run = question_para.add_run("分析質問: ")
        question_run.bold = True
        question_para.add_run(section.question)

        # 内容
        self._doc.add_paragraph(section.content)

        # グラフ画像
        for chart_path in section.chart_paths:
            if Path(chart_path).exists():
                try:
                    self._doc.add_picture(chart_path, width=Inches(5))
                    # キャプション
                    caption_para = self._doc.add_paragraph()
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_run = caption_para.add_run(f"図 {index}: 分析結果グラフ")
                    caption_run.font.size = Pt(10)
                    caption_run.font.italic = True
                except Exception as e:
                    logger.warning(f"Failed to add chart image: {e}")

        self._doc.add_paragraph()

    def _add_conclusion_section(self, conclusion: str) -> None:
        """結論セクションを追加"""
        if self._doc is None:
            return

        self._doc.add_heading("結論", level=1)
        self._doc.add_paragraph(conclusion)

    def _save(self) -> str:
        """文書を一時ファイルに保存"""
        if self._doc is None:
            raise ValueError("Document not initialized")

        # 一時ファイルとして保存
        tmp_dir = tempfile.gettempdir()
        filename = f"report_{uuid.uuid4().hex[:8]}.docx"
        local_path = str(Path(tmp_dir) / filename)

        self._doc.save(local_path)
        return local_path
